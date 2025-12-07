import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
import re


def convert_markdown_to_docx(md_content: str, output_path: str):
    """
    Convert markdown content to Word document with basic formatting.

    Args:
        md_content: Markdown content to convert
        output_path: Path to save the .docx file
    """
    doc = Document()

    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            bullet_items = []
            while i < len(lines) and (lines[i].startswith('- ') or lines[i].startswith('* ')):
                bullet_items.append(lines[i][2:])
                i += 1
            for item in bullet_items:
                doc.add_paragraph(item, style='List Bullet')
            continue
        elif re.match(r'^\d+\.\s', line):
            numbered_items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i]):
                numbered_items.append(re.sub(r'^\d+\.\s', '', lines[i]))
                i += 1
            for item in numbered_items:
                doc.add_paragraph(item, style='List Number')
            continue
        elif line.strip():
            doc.add_paragraph(line)

        i += 1

    doc.save(output_path)


def save_markdown_document(content: str, project_name: str, doc_type: str, output_dir: str = "./output") -> dict:
    """
    Save markdown content to both .md and .docx files with Spanish naming convention.

    Args:
        content: Markdown content to save
        project_name: Name of the project
        doc_type: Type of document (REQUERIMIENTOS, PDD, HISTORIAS_USUARIO)
        output_dir: Directory to save the document

    Returns:
        Dictionary with paths to both files: {"md": md_path, "docx": docx_path}
    """
    os.makedirs(output_dir, exist_ok=True)

    nombre = project_name.lower().replace(" ", "-").replace("_", "-")
    filename_base = f"{doc_type}-{nombre}"

    md_filepath = os.path.join(output_dir, f"{filename_base}.md")
    docx_filepath = os.path.join(output_dir, f"{filename_base}.docx")

    with open(md_filepath, "w", encoding="utf-8") as f:
        f.write(content)

    convert_markdown_to_docx(content, docx_filepath)

    return {
        "md": md_filepath,
        "docx": docx_filepath
    }


def load_document(filepath: str) -> str:
    """
    Load document content from file.

    Args:
        filepath: Path to the document

    Returns:
        Document content as string
    """
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()
